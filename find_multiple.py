from docx import Document
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc = Document(r"E:\临时\思想道德与法治习题库(1).docx")

print("=== Looking for multiple choice section ===")

for i, para in enumerate(doc.paragraphs[50:200]):
    text = para.text.strip()
    if text:
        print(f"[{i+50}] {repr(text)[:150]}")
        if '多选' in text:
            print("======= FOUND MULTIPLE SECTION =======")
            for j in range(1, 30):
                if i+50+j < len(doc.paragraphs):
                    t = doc.paragraphs[i+50+j].text.strip()
                    if t:
                        print(f"  [{i+50+j}] {repr(t)[:150]}")
            break
