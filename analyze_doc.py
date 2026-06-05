from docx import Document
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc = Document(r"E:\临时\思想道德与法治习题库(1).docx")

print("=== Document structure analysis ===")
print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")

print("\n=== First 50 paragraphs ===")
for i, para in enumerate(doc.paragraphs[:50]):
    text = para.text.strip()
    if text:
        print(f"[{i}] {repr(text)[:100]}")
        if any(run.bold for run in para.runs):
            print(f"    -> BOLD")

print("\n=== Tables ===")
for table_idx, table in enumerate(doc.tables[:5]):
    print(f"\nTable {table_idx}: {len(table.rows)} rows, {len(table.columns)} columns")
    for row_idx, row in enumerate(table.rows[:3]):
        row_text = [cell.text.strip() for cell in row.cells]
        print(f"  Row {row_idx}: {row_text}")
